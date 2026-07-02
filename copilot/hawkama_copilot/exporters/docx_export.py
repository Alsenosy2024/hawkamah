"""DOCX export with full Arabic RTL support (python-docx).

python-docx has no high-level RTL switch, so we set the low-level OOXML bits
directly: document-default bidi, per-paragraph `w:bidi`, right alignment, and the
`w:rtl` run property. Default body font is Almarai (the project mandate for Word
exports — do NOT swap to Thmanyah here). Tables get a navy header + zebra rows.

The Almarai TTFs are also EMBEDDED into the fontTable so the document still
renders in the brand face on a machine that doesn't have Almarai installed
(mirrors the PDF exporter's vendored-font approach) — see ``_embed_font``.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.oxml import serialize_part_xml
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import nsmap as OOXML_NS, qn
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from lxml import etree

from .markdown_ast import Block, parse_markdown

NAVY = RGBColor(0x1A, 0x35, 0x57)
GOLD = RGBColor(0xC8, 0x91, 0x2A)
DEFAULT_FONT = "Almarai"

# Vendored inside the package (present in the container image), mirroring
# pdf_export._PKG_FONTS.
_FONT_DIR = Path(__file__).resolve().parents[1] / "fonts"
# The OOXML spec's documented signal that an embedded font's bytes are stored
# AS-IS (unobfuscated): Word reads an `application/x-font-ttf` font part as
# plain TrueType data when its `w:fontKey` is this all-zero GUID. Real Word GUID
# -XOR obfuscation is an optional, purely cosmetic anti-copy measure — not
# required for the embed itself to work — so we skip re-implementing that
# byte-shuffling algorithm here and use the simpler, equally spec-valid path.
_UNOBFUSCATED_FONT_KEY = "00000000-0000-0000-0000-000000000000"


def export_docx(markdown: str, title: str, *, font: str = DEFAULT_FONT, company: str = "") -> bytes:
    doc = Document()
    _set_rtl_defaults(doc, font)
    _cover(doc, title, company, font)

    for b in parse_markdown(markdown):
        _emit_block(doc, b, font)

    _embed_font(doc, font)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Font embedding (D5)                                                          #
# --------------------------------------------------------------------------- #
def _embed_font(doc: Document, font_name: str) -> None:
    """Embed the Almarai TTF (regular + bold) into the docx's fontTable.

    python-docx has no high-level font-embedding API, so this builds the OOXML
    pieces directly: font-data parts for the TTFs, `<w:font>` entries in the
    document's fontTable part referencing them, and the `w:embedTrueTypeFonts`
    settings.xml flag that tells Word to actually use the embedded bytes rather
    than substituting a locally-installed font.

    python-docx's blank-document template already SHIPS an (unembedded) fontTable
    part listing the theme fonts (Calibri, Cambria, …) — we reuse that existing
    part rather than adding a second, competing one at the same partname. A
    missing vendored TTF, or a non-Almarai `font_name`, is a silent no-op: the
    document still renders, just without the embed (exactly the prior behavior).
    """
    if font_name != DEFAULT_FONT:
        return  # only the project's mandated Almarai face is vendored
    regular_path = _FONT_DIR / "Almarai-Regular.ttf"
    if not regular_path.is_file():
        return

    package = doc.part.package
    document_part = doc.part

    try:
        font_table_part = document_part.part_related_by(RT.FONT_TABLE)
        root = etree.fromstring(font_table_part.blob)
    except (KeyError, ValueError):
        font_table_part = Part(PackURI("/word/fontTable.xml"), CT.WML_FONT_TABLE, b"", package)
        document_part.relate_to(font_table_part, RT.FONT_TABLE)
        root = etree.Element(qn("w:fonts"), nsmap={"w": OOXML_NS["w"], "r": OOXML_NS["r"]})

    font_el = etree.SubElement(root, qn("w:font"))
    font_el.set(qn("w:name"), font_name)

    def embed(tag: str, path: Path) -> None:
        if not path.is_file():
            return
        # `next_partname` walks the package's relationship graph, so each font
        # data part must be related in BEFORE the next one is named, or two
        # unrelated (not-yet-linked) parts both compute the same "next" name.
        data_part = Part(
            package.next_partname("/word/fonts/font%d.fntdata"), CT.X_FONT_TTF,
            path.read_bytes(), package,
        )
        rid = font_table_part.relate_to(data_part, RT.FONT)
        el = etree.SubElement(font_el, qn(tag))
        el.set(qn("r:id"), rid)
        el.set(qn("w:fontKey"), _UNOBFUSCATED_FONT_KEY)

    embed("w:embedRegular", regular_path)
    embed("w:embedBold", _FONT_DIR / "Almarai-Bold.ttf")
    font_table_part._blob = serialize_part_xml(root)  # Part.blob has no public setter

    # CT_Settings element order: embedTrueTypeFonts sits between printFormsData
    # and embedSystemFonts — both absent from python-docx's template — so it must
    # come right after w:zoom (the schema only constrains the relative order of
    # PRESENT siblings).
    settings_el = doc.settings.element
    embed_flag = OxmlElement("w:embedTrueTypeFonts")
    embed_flag.set(qn("w:val"), "true")
    zoom_el = settings_el.find(qn("w:zoom"))
    if zoom_el is not None:
        zoom_el.addnext(embed_flag)
    else:
        settings_el.insert(0, embed_flag)


# --------------------------------------------------------------------------- #
def _set_rtl_defaults(doc: Document, font: str) -> None:
    # Default font.
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), font)   # complex-script font = Arabic
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    # Document default bidi.
    _bidi_ppr(style.element.get_or_add_pPr())

    # Section direction: make every section RTL.
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = sectPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            sectPr.append(bidi)


def _bidi_ppr(ppr) -> None:
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))


def _rtl_para(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _bidi_ppr(p._p.get_or_add_pPr())
    for run in p.runs:
        rpr = run._element.get_or_add_rPr()
        if rpr.find(qn("w:rtl")) is None:
            rpr.append(OxmlElement("w:rtl"))


def _cover(doc: Document, title: str, company: str, font: str) -> None:
    h = doc.add_paragraph()
    run = h.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.name = font
    _rtl_para(h)
    if company:
        c = doc.add_paragraph()
        cr = c.add_run(company)
        cr.font.size = Pt(14)
        cr.font.color.rgb = GOLD
        cr.font.name = font
        _rtl_para(c)
    doc.add_page_break()


def _emit_block(doc: Document, b: Block, font: str) -> None:
    if b.type == "heading":
        p = doc.add_heading(level=min(max(b.level, 1), 4))
        p.clear()
        run = p.add_run(b.text)
        run.font.name = font
        run.font.color.rgb = NAVY
        _rtl_para(p)
    elif b.type == "paragraph":
        p = doc.add_paragraph(b.text)
        _style_runs(p, font)
        _rtl_para(p)
    elif b.type == "bullet":
        p = doc.add_paragraph(b.text, style="List Bullet")
        _style_runs(p, font)
        _rtl_para(p)
    elif b.type == "ordered":
        p = doc.add_paragraph(b.text, style="List Number")
        _style_runs(p, font)
        _rtl_para(p)
    elif b.type == "quote":
        p = doc.add_paragraph(b.text, style="Intense Quote")
        _style_runs(p, font)
        _rtl_para(p)
    elif b.type in ("code", "mermaid"):
        # Word has no Mermaid renderer; keep the diagram source as monospace text so
        # it is preserved rather than dropped (rich rendering lives in the HTML export).
        p = doc.add_paragraph(b.text)
        for run in p.runs:
            run.font.name = "Consolas"
    elif b.type == "rule":
        doc.add_paragraph("―" * 20)
    elif b.type == "table":
        _emit_table(doc, b, font)


def _style_runs(p, font: str) -> None:
    for run in p.runs:
        run.font.name = font


def _emit_table(doc: Document, b: Block, font: str) -> None:
    cols = max(len(b.headers), max((len(r) for r in b.rows), default=0))
    if cols == 0:
        return
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _table_rtl(table)

    hdr = table.rows[0].cells
    for i in range(cols):
        text = b.headers[i] if i < len(b.headers) else ""
        cell = hdr[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.bold = True
        run.font.name = font
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, "1A3557")
        _rtl_para(cell.paragraphs[0])

    for r in b.rows:
        cells = table.add_row().cells
        for i in range(cols):
            text = r[i] if i < len(r) else ""
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(text)
            run.font.name = font
            _rtl_para(cells[i].paragraphs[0])


def _table_rtl(table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))


def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
